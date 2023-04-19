import docx
with open("Input/Letters/starting_letter.docx") as starting:
    content= starting.read()


with open("Input/Names/invited_names.txt") as names:
    for name in names:
        if name == "\n":
            continue
        elif name[-1]=="\n":
            name=name[0:-1]
            text = content.replace("[name]", name)
            file = docx.Document()
            file.add_paragraph(text)
            file.save(str("Output/ReadyToSend/letter_for_" + name + ".docx"))
        else:
            text = content.replace("[name]", name)
            file = docx.Document()
            file.add_paragraph(text)
            file.save(str("Output/ReadyToSend/letter_for_" + name+".docx"))



